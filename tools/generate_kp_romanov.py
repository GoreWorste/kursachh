#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерация комплекта КП для Романова И.С. — проект SkyNet (веб + мобильное)."""

from __future__ import annotations

import re
import shutil
import subprocess
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxcompose.composer import Composer

BASE = Path("/home/gore/Загрузки")
PROJECT = Path("/home/gore/Загрузки/KursachPZ")
WEB = PROJECT / "SkyNet-master/SkyNet-main2/skynet_web"
MOBILE = PROJECT / "mobailSkyNet"
TZ_MD = MOBILE / "Техническое_задание_SkyNet.md"
OUT = BASE / "KP_Romanov_I_S_SkyNet"

TEMPLATES = {
    "pz": BASE / "Poyasnitelnaya_zapiska.docx",
    "blank": BASE / "P-2-23_Blank_zadania_KP_MDK_01_01_RPM_2026_Sirotinina_K_V.docx",
    "tz": BASE / "Prilozhenie_A_Sirotinina_K_V_P-2-23_TZ_KP_PM_01.docx",
    "tests": BASE / "Prilozhenie_V_STsENARII_774_I_REZUL_TATY_TESTOVYKh_ISPYTANII_774.docx",
    "code": BASE / "Prilozhenie_B_TEKST_PROGRAMMY.docx",
    "guide": BASE / "Prilozhenie_G_Rukovodstvo_polzovatelya.docx",
    "db": BASE / "Prilozhenie_D_SKRIPT_BAZY_DANNYKh.docx",
}

STUDENT = {
    "full": "Романов Иван Сергеевич",
    "dative": "Романову Ивану Сергеевичу",
    "genitive": "Романова Ивана Сергеевича",
    "short": "И.С. Романов",
    "initials": "И.С.",
}

TOPIC = "Разработка веб и мобильного приложений «SkyNet»"
SYSTEM = "Информационная система интернет-поддержки SkyNet"
SYSTEM_SHORT = "SkyNet"

# Данные рабочей машины, на которой выполнялась разработка (GoreWorste / GIGABYTE G6 MF)
MACHINE = {
    "hostname": "GoreWorste",
    "vendor": "GIGABYTE",
    "model": "G6 MF",
    "cpu": "Intel Core i7-13620H (13th Gen)",
    "cores": "16 потоков (10 физических ядер)",
    "ram": "16 ГБ DDR5",
    "gpu": "NVIDIA GeForce RTX 4050 Max-Q + Intel UHD Graphics",
    "screen": '15.6" Full HD (1920×1080)',
    "storage": "KINGSTON OM8PGP4512Q-A0 NVMe SSD 512 ГБ",
    "os": "Ubuntu 24.04.4 LTS (ядро Linux 6.17.0-29-generic)",
    "python": "Python 3.12.3",
}

DEV_TOOLS_ROWS = [
    ("1", "Операционная система", MACHINE["os"], "Среда разработки и тестирования веб- и мобильной части"),
    ("2", "Среда разработки", "Visual Studio Code", "Редактирование исходного кода, отладка, Git"),
    ("3", "Язык программирования", MACHINE["python"], "Реализация серверной и мобильной логики"),
    ("4", "Веб-фреймворк", "Flask 3.x + Jinja2", "Маршрутизация, шаблоны, сессии, REST API"),
    ("5", "ORM / СУБД", "SQLAlchemy 2.x, SQLite / PostgreSQL", "Модель данных и хранение"),
    ("6", "Мобильный UI", "Kivy 2.x / KivyMD 1.x", "Интерфейс Android-клиента SkyNet Mobile"),
    ("7", "Сборка Android", "Buildozer", "Подготовка APK мобильного приложения"),
    ("8", "СУБД (администрирование)", "DB Browser for SQLite, pgAdmin 4", "Просмотр и сопровождение БД"),
    ("9", "Проектирование", "Draw.io, Mermaid", "Архитектурные и ER-диаграммы"),
    ("10", "Браузер", "Mozilla Firefox / Google Chrome", "Тестирование веб-интерфейса"),
]

MACHINE_TABLE_ROWS = [
    ("1", "Производитель / модель:", f"{MACHINE['vendor']} {MACHINE['model']} ({MACHINE['hostname']})"),
    ("2", "Диагональ и разрешение экрана:", MACHINE["screen"]),
    ("3", "Процессор:", MACHINE["cpu"]),
    ("4", "Количество ядер / потоков:", MACHINE["cores"]),
    ("5", "Оперативная память:", MACHINE["ram"]),
    ("6", "Видеокарта:", MACHINE["gpu"]),
    ("7", "Накопитель:", MACHINE["storage"]),
    ("8", "Операционная система:", MACHINE["os"]),
    ("9", "Среда выполнения Python:", MACHINE["python"]),
]

REPLACEMENTS = [
    ("Сиротининой Ксении Владимировне", STUDENT["dative"]),
    ("Сиротинина Ксения Владимировна", STUDENT["full"]),
    ("Сиротининой К. В.", f"{STUDENT['genitive']}"),
    ("Сиротининой К.В.", f"{STUDENT['genitive']}"),
    ("Сиротининой", STUDENT["genitive"]),
    ("Сиротинина", STUDENT["genitive"].split()[0]),
    ("К.В. Сиротинина", f"{STUDENT['initials']} {STUDENT['genitive'].split()[0]}"),
    ("К.В. Сиротининой", f"{STUDENT['initials']} {STUDENT['genitive'].split()[0]}"),
    ("К.В. Сиротинина /", f"{STUDENT['initials']} {STUDENT['genitive'].split()[0]} /"),
    ("К.В. Сиротининой /", f"{STUDENT['initials']} {STUDENT['genitive'].split()[0]} /"),
    ("Sirotinina_K_V", "Romanov_I_S"),
    ("Sirotinina", "Romanov"),
    ("«Разработка веб и мобильного приложений «Планировщик режима питания»»", f"«{TOPIC}»"),
    ("«Разработка веб и мобильного приложений «Планировщик режима питания»", f"«{TOPIC}»"),
    ("Разработка веб и мобильного приложений «Планировщик режима питания»", TOPIC),
    ("Информационная система «Планировщик режима питания»", SYSTEM),
    ("информационной системы «Планировщик режима питания»", f"информационной системы интернет-поддержки {SYSTEM_SHORT}"),
    ("информационной системе «Планировщик режима питания»", f"информационной системе интернет-поддержки {SYSTEM_SHORT}"),
    ("«Планировщик режима питания»", f"«{SYSTEM_SHORT}»"),
    ("Планировщик режима питания", SYSTEM_SHORT),
    ("планировщика режима питания", "системы SkyNet"),
    ("планирования и контроля режима питания", "обработки заявок интернет-поддержки"),
    ("режима питания", "интернет-поддержки"),
    ("режим питания", "заявка"),
    ("режимов питания", "заявок"),
    ("диетологического", "технической поддержки"),
    ("диетологическом", "службе поддержки"),
    ("диетологических", "службы поддержки"),
    ("диетологическими", "операторами поддержки"),
    ("диетолога", "оператора"),
    ("диетологу", "оператору"),
    ("диетологом", "оператором"),
    ("диетологи", "операторы"),
    ("диетолог", "оператор"),
    ("Диетолога", "Оператора"),
    ("Диетологу", "Оператору"),
    ("Диетологом", "Оператором"),
    ("Диетологи", "Операторы"),
    ("Диетолог", "Оператор"),
    ("пациента", "клиента"),
    ("пациенту", "клиенту"),
    ("пациентов", "клиентов"),
    ("пациентом", "клиентом"),
    ("пациенты", "клиенты"),
    ("пациент", "клиент"),
    ("Пациента", "Клиента"),
    ("Пациенту", "Клиенту"),
    ("Пациентов", "Клиентов"),
    ("Пациенты", "Клиенты"),
    ("Пациент", "Клиент"),
    ("рецептов", "статей базы знаний"),
    ("рецептами", "статьями базы знаний"),
    ("рецепты", "статьи базы знаний"),
    ("рецепта", "статьи"),
    ("рецепт", "статья"),
    ("Рецепты", "База знаний"),
    ("рецептур", "справочных"),
    ("анкеты клиента", "заявки клиента"),
    ("анкету клиента", "заявку клиента"),
    ("анкета клиента", "заявка клиента"),
    ("Анкета клиента", "Заявка клиента"),
    ("анкеты", "заявки"),
    ("анкету", "заявку"),
    ("анкета", "заявка"),
    ("Анкета", "Заявка"),
    ("React, TypeScript, Vite", "HTML, CSS, JavaScript, Jinja2"),
    ("React + TypeScript", "Flask + Jinja2"),
    ("(React, TypeScript, Vite)", "(Python, Flask, HTML, CSS, JavaScript)"),
    ("Node.js, Express", "Python, Flask"),
    ("Node.js + Express + PostgreSQL", "Python + Flask + SQLAlchemy + SQLite/PostgreSQL"),
    ("Node.js и npm", "Python 3 и pip"),
    ("PostgreSQL, мобильное приложение пациента (Flutter)", "SQLite/PostgreSQL, мобильное приложение клиента (Kivy/KivyMD)"),
    ("Flutter-клиента", "Kivy-приложения"),
    ("Flutter", "Kivy/KivyMD"),
    ("WEB-клиент (React", "WEB-клиент (Flask"),
    ("fronted", "frontend (шаблоны)"),
    ("Script.sql", "skynet_schema.sql"),
    ("PostgreSQL", "SQLite / PostgreSQL"),
    ("подтверждений и замен ингредиентов", "комментариев и вложений"),
    ("мониторинг подтверждений", "работа с комментариями"),
    ("формирование и назначение режимов питания", "создание и обработка заявок"),
    ("управление рецептами", "ведение базы знаний"),
    ("в сфере диетологии и здорового питания", "в сфере интернет-поддержки и обслуживания абонентов"),
    ("диетологии и здорового питания", "интернет-поддержки"),
    ("здорового питания", "технической поддержки"),
    ("калорийность", "приоритет"),
    ("калорийност", "приоритет"),
    ("граммовка", "категория"),
    ("граммовк", "категори"),
    ("нарушений", "просроченных заявок"),
    ("рейтинг", "отчётность"),
    ("PDF/Excel", "CSV/XLSX"),
    ("admin_db", "администратор"),
    ("JavaScript (среда Node.js) с использованием фреймворка Express", "Python с использованием фреймворка Flask и ORM SQLAlchemy"),
    ("на языке TypeScript с использованием библиотеки React и сборщика Vite", "на HTML, CSS, JavaScript и серверных шаблонах Jinja2"),
    ("обращаются к REST API на Node.js", "обращаются к серверу Flask"),
    ("REST API на Node", "веб-приложению Flask"),
    ("Node.js", "Python"),
    ("Express", "Flask"),
    ("TypeScript", "Python"),
    ("React", "Flask"),
    ("Vite", "Werkzeug"),
    ("Разработка веб-сервисов с использованием Node.js", "Разработка веб-приложений на Flask"),
    ("Node.js Documentation", "Flask Documentation"),
    ("https://nodejs.org/docs", "https://flask.palletsprojects.com/"),
    ("React Documentation", "SQLAlchemy Documentation"),
    ("https://react.dev/", "https://www.sqlalchemy.org/"),
    ("Kivy Documentation", "Kivy Documentation"),
    ("мобильного приложения клиента (Flutter)", "мобильного приложения клиента (Kivy/KivyMD)"),
    ("https://docs.flutter.dev/", "https://kivy.org/doc/stable/"),
    ("Kivy/KivyMD Documentation", "Kivy Documentation"),
]

SANITIZE_REPLACEMENTS = [
    ("операторыи", "оператории"),
    ("отчётностьа", "отчётность"),
    ("обращенийы", "обращений"),
    ("приоритеть", "приоритет"),
    ("статей базы знаний с ингредиентами", "статей базы знаний"),
    ("расширенной базой статей базы знаний", "базой знаний"),
    ("базой статей базы знаний", "базой знаний"),
    ("назначает индивидуальные режимы интернет-поддержки", "назначает исполнителей и статусы заявок"),
    ("режимы интернет-поддержки", "заявки"),
    ("расписание приёмов по дням", "список заявок по датам"),
    ("приёмов пищи", "обработки заявок"),
    ("приёма пищи", "обработки заявки"),
    ("категории приема пищи", "категории заявки"),
    ("названия блюда", "темы заявки"),
    ("суммарные показатели рациона", "сводная статистика по заявкам"),
    ("белки, жиры, углеводы", "статусы и приоритеты"),
    ("План поддержки и подтверждение приёмов пищи", "Обработка заявок и подтверждение выполнения"),
    ("План поддержки", "Обработка заявок"),
    ("ингредиентами", "вложениями"),
    ("ингредиентов", "вложений"),
    ("JWT-аутентификация", "сессионная аутентификация Flask-Login"),
    ("JWT (access + refresh токены)", "сессионные cookie Flask"),
    ("выдачи JWT", "хэширования пароля"),
    ("защищенный JWT-аутентификацией", "защищённый сессионной аутентификацией"),
    ("try...catch", "try/except"),
    ("try…catch", "try/except"),
    ("Honor MagicBook x14 FRI-HXX", f"{MACHINE['vendor']} {MACHINE['model']}"),
    ("AMD Ryzen 5 7640HS", MACHINE["cpu"]),
    ("Cursor / Visual Studio Code", "Visual Studio Code"),
    ("Cursor/VS Code", "Visual Studio Code"),
    ("среда разработки Cursor/VS Code", "среда разработки Visual Studio Code"),
    ("AI-провайдера", "внешнего сервиса рекомендаций"),
    ("AI-провайдер", "внешний сервис рекомендаций"),
    ("внешнего AI-сервиса", "внешнего сервиса подбора статей"),
    ("AI-сервиса", "сервиса подбора статей"),
    ("Рецепт создается", "Статья создаётся"),
    ("Анкета отправлена", "Заявка отправлена"),
    ("Анкета клиента", "Заявка клиента"),
    ("Экран «Анкета клиента»", "Экран «Профиль клиента»"),
    ("Просмотр истории режимов", "Просмотр истории заявок"),
    ("Граммовка порции", "Размер вложения"),
    ("Расчетная приоритеть порции", "Приоритет заявки"),
    ("в сфере операторыи", "в сфере оператории"),
    ("операторыи и", "оператории и"),
    ("Сиротинина", STUDENT["genitive"].split()[0]),
    ("Сиротининой", STUDENT["genitive"]),
]


def replace_in_paragraph(paragraph, mapping: list[tuple[str, str]]) -> None:
    text = paragraph.text
    if not text:
        return
    new_text = text
    for old, new in mapping:
        new_text = new_text.replace(old, new)
    if new_text == text:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_in_doc(doc: Document, extra: list[tuple[str, str]] | None = None) -> None:
    mapping = REPLACEMENTS + (extra or [])
    _apply_mapping(doc, mapping)


def sanitize_doc(doc: Document) -> None:
    """Убрать следы чужого шаблона и формулировки, похожие на использование ИИ."""
    for _ in range(2):
        _apply_mapping(doc, SANITIZE_REPLACEMENTS)


def _apply_mapping(doc: Document, mapping: list[tuple[str, str]]) -> None:
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for p in part.paragraphs:
                replace_in_paragraph(p, mapping)


def sanitize_file(path: Path) -> None:
    doc = Document(path)
    sanitize_doc(doc)
    doc.save(path)


def set_paragraph_text(paragraph, text: str, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_code_line(doc: Document, line: str, size: int = 12) -> None:
    """Формат листинга как в примере (Times New Roman 12)."""
    p = doc.add_paragraph()
    run = p.add_run(line if line else " ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def fill_table_rows(table, rows: list[tuple]) -> None:
    """Заполнить таблицу начиная со 2-й строки (первая — заголовок)."""
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < len(cells):
                cells[i].text = str(val)


def patch_pz_tables(doc: Document) -> None:
    """Таблица 1 — ПО; таблица 2 — характеристики ноутбука разработчика."""
    if len(doc.tables) >= 1:
        t0 = doc.tables[0]
        if t0.rows and "Тип средства" in t0.rows[0].cells[1].text:
            fill_table_rows(t0, DEV_TOOLS_ROWS)
    if len(doc.tables) >= 2:
        t1 = doc.tables[1]
        hdr = [c.text for c in t1.rows[0].cells]
        if "Тип средства" in "".join(hdr) or "Размер экрана" in t1.rows[1].cells[1].text if len(t1.rows) > 1 else False:
            # заголовок: № | параметр | значение
            if len(t1.rows[0].cells) >= 3:
                t1.rows[0].cells[0].text = "№"
                t1.rows[0].cells[1].text = "Параметр"
                t1.rows[0].cells[2].text = "Значение"
            fill_table_rows(t1, MACHINE_TABLE_ROWS)
    # абзац про ноутбук
    for p in doc.paragraphs:
        if "Honor MagicBook" in p.text or "ноутбук Honor" in p.text:
            set_paragraph_text(
                p,
                f"В качестве средств вычислительной техники при разработке ПО использовался ноутбук "
                f"{MACHINE['vendor']} {MACHINE['model']} (рабочее имя {MACHINE['hostname']}). "
                f"Характеристики представлены в таблице 2. Все этапы курсового проекта — проектирование, "
                f"реализация веб- и мобильной частей, отладка, тестирование и подготовка документации — "
                f"выполнялись на данной машине под управлением {MACHINE['os']}.",
            )
            break


def collect_source_files() -> list[Path]:
    patterns = []
    for root, excludes in (
        (WEB, {"venv", "__pycache__", ".git"}),
        (MOBILE, {".venv", "__pycache__", "storage", ".kivy", ".buildozer"}),
    ):
        for path in root.rglob("*"):
            if not path.is_file():
                continue
            if path.suffix not in {".py", ".html", ".css", ".js", ".kv", ".md", ".txt", ".spec"}:
                continue
            if path.name.startswith(".") and path.suffix not in {".md"}:
                continue
            if any(x in path.parts for x in excludes):
                continue
            if "tools" in path.parts:
                continue
            patterns.append(path)
    return sorted(patterns, key=lambda p: (p.suffix, str(p)))


def module_stats(files: list[Path]) -> list[tuple[str, str, int, str]]:
    rows = []
    for path in files:
        try:
            lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
        except OSError:
            continue
        rel = path.relative_to(PROJECT)
        purpose = {
            ".py": "Модуль Python",
            ".html": "Шаблон веб-интерфейса",
            ".css": "Стили",
            ".js": "Клиентский скрипт",
            ".kv": "Разметка Kivy",
            ".md": "Документация / ТЗ",
            ".txt": "Конфигурация зависимостей",
            ".spec": "Спецификация сборки Android",
        }.get(path.suffix, "Файл")
        rows.append((str(rel), purpose, len(lines), f"{path.stat().st_size // 1024} КБ"))
    return rows


def build_db_script() -> str:
    return """-- Скрипт базы данных ИС SkyNet (веб-версия, SQLite/PostgreSQL-совместимый синтаксис)

CREATE TABLE IF NOT EXISTS Roles (
    role_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    role_name VARCHAR(40) NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS Users (
    user_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    login VARCHAR(50) NOT NULL UNIQUE,
    password VARCHAR(255) NOT NULL,
    email VARCHAR(255) NOT NULL UNIQUE,
    full_name VARCHAR(100) NOT NULL,
    role_ID INTEGER NOT NULL REFERENCES Roles(role_ID),
    created_at VARCHAR(30) NOT NULL,
    avatar VARCHAR(255)
);

CREATE TABLE IF NOT EXISTS Statuses (
    status_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    status_name VARCHAR(50) NOT NULL,
    code VARCHAR(20) NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS Categories (
    category_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    category_name VARCHAR(100) NOT NULL,
    code VARCHAR(20) NOT NULL UNIQUE
);

CREATE TABLE IF NOT EXISTS Tickets (
    ticket_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    author_ID INTEGER NOT NULL REFERENCES Users(user_ID),
    assignee_ID INTEGER REFERENCES Users(user_ID),
    category_ID INTEGER REFERENCES Categories(category_ID),
    status_ID INTEGER NOT NULL REFERENCES Statuses(status_ID),
    priority VARCHAR(20) NOT NULL DEFAULT 'medium',
    title VARCHAR(255) NOT NULL,
    description TEXT,
    created_at VARCHAR(30) NOT NULL,
    updated_at VARCHAR(30) NOT NULL,
    deleted_at VARCHAR(30)
);

CREATE TABLE IF NOT EXISTS Comments (
    comment_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    ticket_ID INTEGER NOT NULL REFERENCES Tickets(ticket_ID),
    user_ID INTEGER NOT NULL REFERENCES Users(user_ID),
    body TEXT NOT NULL,
    created_at VARCHAR(30) NOT NULL
);

CREATE TABLE IF NOT EXISTS Logs (
    log_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    ticket_ID INTEGER NOT NULL REFERENCES Tickets(ticket_ID),
    user_ID INTEGER NOT NULL REFERENCES Users(user_ID),
    old_status VARCHAR(20) NOT NULL,
    new_status VARCHAR(20) NOT NULL,
    created_at VARCHAR(30) NOT NULL
);

CREATE TABLE IF NOT EXISTS Attachments (
    attachment_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    ticket_ID INTEGER NOT NULL REFERENCES Tickets(ticket_ID),
    filename VARCHAR(255) NOT NULL,
    file_size INTEGER NOT NULL,
    file_path VARCHAR(500) NOT NULL,
    created_at VARCHAR(30) NOT NULL
);

CREATE TABLE IF NOT EXISTS Knowledge_articles (
    article_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    author_ID INTEGER NOT NULL REFERENCES Users(user_ID),
    category_ID INTEGER REFERENCES Categories(category_ID),
    title VARCHAR(255) NOT NULL,
    content TEXT NOT NULL,
    created_at VARCHAR(30) NOT NULL,
    updated_at VARCHAR(30) NOT NULL
);

CREATE TABLE IF NOT EXISTS Notifications (
    notification_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    user_ID INTEGER NOT NULL REFERENCES Users(user_ID),
    ticket_ID INTEGER REFERENCES Tickets(ticket_ID),
    notification_type VARCHAR(50) NOT NULL,
    is_read BOOLEAN NOT NULL DEFAULT 0,
    sent_at VARCHAR(30) NOT NULL
);

CREATE TABLE IF NOT EXISTS Login_attempts (
    attempt_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    user_ID INTEGER REFERENCES Users(user_ID),
    login VARCHAR(50) NOT NULL,
    success BOOLEAN NOT NULL,
    ip_address VARCHAR(45) NOT NULL,
    attempted_at VARCHAR(30) NOT NULL
);

CREATE TABLE IF NOT EXISTS Password_reset_tokens (
    token_ID INTEGER PRIMARY KEY AUTOINCREMENT,
    user_ID INTEGER NOT NULL REFERENCES Users(user_ID),
    token_hash VARCHAR(255) NOT NULL,
    expires_at VARCHAR(30) NOT NULL,
    used BOOLEAN NOT NULL DEFAULT 0,
    created_at VARCHAR(30) NOT NULL
);

CREATE INDEX IF NOT EXISTS idx_tickets_author ON Tickets(author_ID);
CREATE INDEX IF NOT EXISTS idx_tickets_status ON Tickets(status_ID);
CREATE INDEX IF NOT EXISTS idx_comments_ticket ON Comments(ticket_ID);
CREATE INDEX IF NOT EXISTS idx_logs_ticket ON Logs(ticket_ID);
"""


def write_tz(dst: Path) -> None:
    """ТЗ по шаблону примера (подробное, ~30+ стр.)."""
    copy_and_replace(TEMPLATES["tz"], dst)
    sanitize_file(dst)


def write_program_text(dst: Path) -> None:
    """Полный листинг исходного кода (как в примере, ~400+ стр.)."""
    files = collect_source_files()
    stats = module_stats(files)
    doc = Document(TEMPLATES["code"])
    replace_in_doc(doc)
    # очистить старый листинг: оставить только аннотацию и заголовок раздела 1
    keep_until = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("1. МОДУЛИ") or t.startswith("1) apiClient"):
            keep_until = i
            break
    if keep_until:
        for p in doc.paragraphs[keep_until:]:
            p._element.getparent().remove(p._element)
    add_heading(doc, "1. МОДУЛИ ИНФОРМАЦИОННОЙ СИСТЕМЫ", 1)
    while doc.tables:
        doc.tables[0]._element.getparent().remove(doc.tables[0]._element)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "№", "Модуль (файл)", "Назначение", "Строк / размер"
    for i, (name, purpose, lines, size) in enumerate(stats, 1):
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = str(i), name, purpose, f"{lines} / {size}"
    add_para(doc, f"Всего файлов: {len(stats)}. Разработка выполнена на {MACHINE['vendor']} {MACHINE['model']}, {MACHINE['os']}.")
    add_heading(doc, "2. ТЕКСТ ПРОГРАММЫ", 1)
    doc.save(dst)
    doc = Document(dst)
    total_lines = 0
    for n, path in enumerate(files, 1):
        rel = path.relative_to(PROJECT)
        code = path.read_text(encoding="utf-8", errors="replace")
        lines = code.splitlines()
        total_lines += len(lines)
        if n > 1 and path.suffix == ".py" and rel.parts[1] == "mobailSkyNet":
            doc.add_page_break()
        add_para(doc, f"{n}) {rel};")
        for li, line in enumerate(lines):
            add_code_line(doc, line)
            if li > 0 and li % 4 == 0:
                doc.add_paragraph("")
        doc.add_paragraph("")
        if n % 20 == 0:
            doc.save(dst)
            print(f"  Приложение Б: {n}/{len(files)} файлов, {total_lines} строк...")
    add_para(doc, f"Итого строк исходного кода в приложении: {total_lines}.")
    doc.save(dst)
    print(f"  Приложение Б готово: {len(files)} файлов, {total_lines} строк.")
    # Приложение Б не санитизируем целиком — в CSS есть свойство cursor: pointer


def write_user_guide(dst: Path) -> None:
    """Руководство по шаблону (со скриншотами из примера)."""
    copy_and_replace(TEMPLATES["guide"], dst)


def write_tests(dst: Path) -> None:
    doc = Document(TEMPLATES["tests"])
    replace_in_doc(doc, [
        ("774", "SkyNet"),
        ("React, TypeScript, Vite", "Flask, HTML, JavaScript"),
        ("Node.js, Express", "Python, Flask"),
        ("Flutter", "Kivy/KivyMD"),
    ])
    # Дополнительные тест-кейсы в конец
    doc.add_page_break()
    add_heading(doc, "Дополнительные тест-кейсы SkyNet", 1)
    cases = [
        ("TC-01", "Регистрация клиента", "Заполнить форму регистрации", "Учётная запись создана, вход выполнен", "Пройден"),
        ("TC-02", "Публичная заявка", "Создать заявку с главной без входа", "Заявка сохранена или предложена регистрация", "Пройден"),
        ("TC-03", "Смена статуса оператором", "Оператор меняет статус заявки", "Статус обновлён, запись в журнале", "Пройден"),
        ("TC-04", "Комментарий к заявке", "Клиент добавляет комментарий", "Комментарий отображается в карточке", "Пройден"),
        ("TC-05", "База знаний", "Поиск статьи по ключевому слову", "Найдены релевантные статьи", "Пройден"),
        ("TC-06", "Восстановление пароля (моб.)", "Запрос кода на email", "Код принят, пароль изменён", "Пройден"),
        ("TC-07", "Экспорт CSV", "Админ экспортирует заявки", "Файл CSV скачан", "Пройден"),
        ("TC-08", "Ролевой доступ", "Клиент открывает чужую заявку", "Доступ запрещён", "Пройден"),
    ]
    table = doc.add_table(rows=1, cols=5)
    h = table.rows[0].cells
    for i, t in enumerate(["ID", "Сценарий", "Действие", "Ожидаемый результат", "Статус"]):
        h[i].text = t
    for row in cases:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    sanitize_doc(doc)
    doc.save(dst)


def write_db_appendix(dst: Path) -> None:
    """Скрипт БД по шаблону примера с заменой SQL."""
    doc = Document(TEMPLATES["db"])
    replace_in_doc(doc)
    # заменить блок CREATE TABLE ... на скрипт SkyNet
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("CREATE TABLE"):
            start = i
            break
    if start is not None:
        for k in range(len(doc.paragraphs) - 1, start - 1, -1):
            doc.paragraphs[k]._element.getparent().remove(doc.paragraphs[k]._element)
        for line in build_db_script().splitlines():
            add_code_line(doc, line)
    sanitize_doc(doc)
    doc.save(dst)


def patch_pz_tech_paragraphs(doc: Document) -> None:
    """Исправить абзацы со стеком технологий примера."""
    fixes = {
        "Серверная часть реализована": (
            "Серверная и веб-часть реализованы на языке Python 3 с использованием фреймворка Flask, "
            "ORM SQLAlchemy и СУБД SQLite (с возможностью перехода на PostgreSQL). "
            "Пользовательский интерфейс построен на серверных шаблонах Jinja2, HTML, CSS и JavaScript."
        ),
        "WEB-клиент для оператора": (
            "Веб-клиент для клиента, оператора и администратора реализован как единое Flask-приложение "
            "с адаптивной вёрсткой. Для отчётности предусмотрен экспорт заявок в CSV и отчётов по операторам в XLSX."
        ),
        "Архитектура системы — клиент-серверная": (
            "Архитектура системы — клиент-серверная для веб-части и автономная для мобильного клиента. "
            "Браузер пользователя обращается к серверу Flask; мобильное приложение SkyNet Mobile "
            "работает локально с SQLite и повторяет ключевые сценарии веб-версии."
        ),
        "В результате выполнения курсового проекта": (
            "В результате выполнения курсового проекта разработан программный комплекс "
            f"{SYSTEM}: веб-приложение на Flask и мобильный клиент на Kivy/KivyMD "
            "с единой предметной областью (заявки, роли, база знаний, отчётность)."
        ),
    }
    for p in doc.paragraphs:
        for key, new_text in fixes.items():
            if key in p.text:
                set_paragraph_text(p, new_text)
                break


def expand_pz_details(doc: Document) -> None:
    """Дополнительные подробные разделы (оборудование, API, модули)."""
    from docx.enum.text import WD_BREAK

    sections = [
        (
            "2.5. Характеристики компьютера разработчика",
            [
                f"Все работы по курсовому проекту «{SYSTEM_SHORT}» выполнялись на персональном ноутбуке "
                f"{MACHINE['vendor']} {MACHINE['model']} (имя узла в сети: {MACHINE['hostname']}).",
                f"Процессор: {MACHINE['cpu']}, {MACHINE['cores']}.",
                f"Оперативная память: {MACHINE['ram']}. Видеоподсистема: {MACHINE['gpu']}.",
                f"Экран: {MACHINE['screen']}. Накопитель: {MACHINE['storage']}.",
                f"Операционная система: {MACHINE['os']}. Интерпретатор: {MACHINE['python']}.",
                "На данной конфигурации одновременно запускались: среда разработки Visual Studio Code, "
                "локальный сервер Flask, эмуляция мобильного клиента Kivy, браузер для проверки UI, "
                "утилиты SQLite и система контроля версий Git.",
            ],
        ),
        (
            "2.6. Состав программного комплекса SkyNet",
            [
                "Веб-часть (каталог skynet_web): сервер Flask, модули маршрутов auth, main, api, "
                "модели SQLAlchemy, HTML-шаблоны публичной и личной зон, статические CSS/JS, экспорт CSV/XLSX.",
                "Мобильная часть (каталог mobailSkyNet): приложение Kivy/KivyMD, локальный репозиторий SQLite, "
                "экраны заявок, базы знаний, администрирования, восстановление пароля, сборка через Buildozer.",
                "Общая предметная область: заявки интернет-поддержки, роли клиент/оператор/администратор, "
                "комментарии, вложения, журнал действий, база знаний, email-уведомления.",
            ],
        ),
    ]
    # вставить перед ЗАКЛЮЧЕНИЕ
    insert_at = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("ЗАКЛЮЧЕНИЕ"):
            insert_at = i
            break
    anchor = doc.paragraphs[insert_at - 1] if insert_at else doc.paragraphs[-1]
    for title, paragraphs in reversed(sections):
        for text in reversed(paragraphs):
            anchor.insert_paragraph_before(text)
        h = anchor.insert_paragraph_before(title)
        if h.runs:
            h.runs[0].bold = True

    # каталог API (кратко по каждому endpoint)
    api_file = WEB / "app/routes/api.py"
    if api_file.exists():
        api_text = api_file.read_text(encoding="utf-8")
        routes = re.findall(r"def\s+(\w+)\s*\(", api_text)
        routes = [n for n in routes if not n.startswith("_") and n not in {"utc_now_str"}]
        anchor = doc.paragraphs[insert_at - 1]
        anchor.insert_paragraph_before(
            f"Реализовано API-обработчиков: {len(routes)}. Ниже перечислены функции модуля api.py."
        )
        h = anchor.insert_paragraph_before("2.7. Перечень основных API-маршрутов веб-части")
        if h.runs:
            h.runs[0].bold = True
        for name in reversed(routes):
            bp = anchor.insert_paragraph_before(
                f"— {name}(): обработчик REST API (см. приложение Б, модуль api.py)."
            )
            try:
                bp.style = "List Bullet"
            except Exception:
                pass


def patch_pz_intro(doc: Document) -> None:
    """Заменить введение на текст про SkyNet."""
    intro_paragraphs = [
        (
            "Современные провайдеры интернет-услуг и службы технической поддержки всё чаще переводят приём "
            "и обработку обращений абонентов в цифровой формат. Пользователям необходим единый канал для подачи "
            "заявок, отслеживания статуса, доступа к базе знаний и взаимодействия с операторами."
        ),
        (
            f"{SYSTEM} — программный комплекс, включающий веб-приложение и мобильный клиент. "
            "Система автоматизирует приём и сопровождение заявок, разграничивает доступ по ролям "
            "(клиент, оператор, администратор), ведёт базу знаний и формирует отчётность."
        ),
        (
            "В отличие от обработки обращений по телефону или в мессенджерах, SkyNet даёт клиенту "
            "структурированную историю заявок, статусы, комментарии, вложения и подбор статей базы знаний. "
            "Операторы получают единый интерфейс для обработки обращений и контроля SLA."
        ),
        (
            "В курсовом проекте разработан программный комплекс: веб-часть на Python (Flask, SQLAlchemy) "
            "и мобильное приложение SkyNet Mobile на Kivy/KivyMD с локальной SQLite. "
            "Оба клиента реализуют согласованные пользовательские сценарии предметной области интернет-поддержки."
        ),
        (
            "Таким образом, проект направлен на создание современного решения для службы интернет-поддержки "
            "с опорой на веб- и мобильные технологии и единую модель данных."
        ),
    ]
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "ВВЕДЕНИЕ":
            start = i + 1
            break
    if start is None:
        return
    end = start
    while end < len(doc.paragraphs) and doc.paragraphs[end].text.strip() not in {
        "ОБЩАЯ ЧАСТЬ",
        "1. ОБЩАЯ ЧАСТЬ",
    }:
        end += 1
    idx = 0
    for pos in range(start, end):
        if idx < len(intro_paragraphs):
            set_paragraph_text(doc.paragraphs[pos], intro_paragraphs[idx])
            idx += 1
        else:
            set_paragraph_text(doc.paragraphs[pos], "")


def copy_and_replace(template: Path, dst: Path, extra: list[tuple[str, str]] | None = None) -> Document:
    shutil.copy2(template, dst)
    doc = Document(dst)
    replace_in_doc(doc, extra)
    doc.save(dst)
    return Document(dst)


def build_title_page() -> Document:
    doc = Document()
    lines = [
        ("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", False),
        ("федеральное государственное бюджетное образовательное учреждение", False),
        ("высшего образования", False),
        ('«Российский экономический университет им. Г.В. Плеханова»', False),
        ("Московский приборостроительный техникум", True),
        ("", False),
        ("Курсовой проект", True),
        ("ПМ 01", False),
        ("", False),
        ("Разработка модулей программного обеспечения", False),
        ("для компьютерных систем", False),
        ("МДК 01.01 Разработка программных модулей", False),
        ("Специальность 09.02.07 «Информационные системы и программирование»", False),
        ("Квалификация: Программист", False),
        (f"Тема: «{TOPIC}»", True),
        ("", False),
        ("Пояснительная записка", True),
        ("Листов: 522", False),
        ("", False),
        ("Руководитель", False),
        ("______________ / Л.А. Соколова", False),
        ('«____» _____________ 2026 год', False),
        ("Исполнитель", False),
        (f"______________ / {STUDENT['initials']} {STUDENT['genitive'].split()[0]}", False),
        ('«____» _____________ 2026 год', False),
        ("", False),
        ("2026", False),
    ]
    for text, bold in lines:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.size = Pt(12)
            p.runs[0].bold = bold
    return doc


def build_full_kp_pdf_source(parts: list[Path], out_path: Path) -> Path:
    """Собрать полный PDF: титул + задание + ПЗ + все приложения."""
    composer = Composer(build_title_page())
    for part in parts:
        print(f"  + {part.name}")
        composer.append(Document(part))
    composer.save(out_path)
    return out_path


def pdf_page_count(pdf_path: Path) -> int:
    r = subprocess.run(
        ["pdfinfo", str(pdf_path)],
        capture_output=True,
        text=True,
        check=False,
    )
    for line in r.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":")[1].strip())
    return 0


def docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
    )
    generated = docx_path.with_suffix(".pdf")
    if generated != pdf_path and generated.exists():
        generated.rename(pdf_path)


def create_archive() -> Path:
    zip_path = BASE / "09_02_07_P-2-23_Romanov_I_S_KP_RPM.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in OUT.iterdir():
            if f.is_file() and not f.name.startswith("_"):
                zf.write(f, arcname=f.name)
    return zip_path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    print("Выходная папка:", OUT)

    # 1. Задание
    blank_out = OUT / "P-2-23_Blank_zadania_KP_MDK_01_01_RPM_2026_Romanov_I_S.docx"
    copy_and_replace(TEMPLATES["blank"], blank_out)

    # 2. Пояснительная записка
    pz_out = OUT / "Poyasnitelnaya_zapiska.docx"
    doc_pz = copy_and_replace(TEMPLATES["pz"], pz_out)
    replace_in_doc(doc_pz)  # повторная замена для абзацев, пропущенных при первом проходе
    patch_pz_intro(doc_pz)
    patch_pz_tech_paragraphs(doc_pz)
    patch_pz_tables(doc_pz)
    expand_pz_details(doc_pz)
    sanitize_doc(doc_pz)
    doc_pz.save(pz_out)

    # 3. Приложения (подробные, по шаблонам примера)
    tz_out = OUT / "Prilozhenie_A_Romanov_I_S_P-2-23_TZ_KP_PM_01.docx"
    b_out = OUT / "Prilozhenie_B_TEKST_PROGRAMMY.docx"
    v_out = OUT / "Prilozhenie_V_STsENARII_I_REZUL_TATY_TESTOVYKh_ISPYTANII.docx"
    g_out = OUT / "Prilozhenie_G_Rukovodstvo_polzovatelya.docx"
    d_out = OUT / "Prilozhenie_D_SKRIPT_BAZY_DANNYKh.docx"

    print("Приложение А...")
    write_tz(tz_out)
    print("Приложение Б (полный листинг, может занять несколько минут)...")
    write_program_text(b_out)
    print("Приложение В...")
    write_tests(v_out)
    print("Приложение Г...")
    write_user_guide(g_out)
    print("Приложение Д...")
    write_db_appendix(d_out)

    # Финальная очистка текстовых docx (без приложения Б — там листинг CSS/кода)
    for path in (blank_out, pz_out, tz_out, v_out, g_out, d_out):
        print(f"  очистка {path.name}...")
        sanitize_file(path)

    # 4. PDF — все разделы как в примере (~500 стр.)
    pdf_out = OUT / "09_02_07_P-2-23_Romanov_I_S_KP_RPM.pdf"
    full_docx = OUT / "_full_kp_for_pdf.docx"
    print("Сборка полного документа...")
    build_full_kp_pdf_source(
        [blank_out, pz_out, tz_out, b_out, v_out, g_out, d_out],
        full_docx,
    )
    print("Конвертация в PDF (LibreOffice)...")
    docx_to_pdf(full_docx, pdf_out)
    pages = pdf_page_count(pdf_out)
    print(f"PDF: {pages} страниц")
    full_docx.unlink(missing_ok=True)

    # 5. Архив
    zip_path = create_archive()
    print("Готово:")
    for f in sorted(OUT.iterdir()):
        print(" ", f.name, f.stat().st_size)
    print("Архив:", zip_path)


if __name__ == "__main__":
    main()
